"""Convert RA markdown to a formatted Word document (.docx)."""
import io
import re

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Inches


def _shade_cell(cell, hex_color: str) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _write_inline(para, text: str) -> None:
    """Add runs for **bold** and plain text. Strips leftover * markers."""
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = para.add_run(part[2:-2])
            run.bold = True
        else:
            para.add_run(part)


def _is_separator_row(row_line: str) -> bool:
    return bool(re.match(r"^\|[\s\-:|]+\|$", row_line.strip()))


def markdown_to_docx(ra_markdown: str, project_name: str = "Risk Assessment") -> bytes:
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)

    # Base font
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10)

    lines = ra_markdown.splitlines()
    i = 0

    while i < len(lines):
        raw = lines[i]
        stripped = raw.strip()

        # Skip horizontal rules
        if re.match(r"^[-*_]{3,}$", stripped):
            i += 1
            continue

        # Headings
        m = re.match(r"^(#{1,4})\s+(.*)", stripped)
        if m:
            level = min(len(m.group(1)), 4)
            doc.add_heading(m.group(2).strip(), level=level)
            i += 1
            continue

        # Table block — collect all pipe lines
        if stripped.startswith("|"):
            raw_rows: list[str] = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                if not _is_separator_row(lines[i]):
                    raw_rows.append(lines[i].strip())
                i += 1

            if not raw_rows:
                continue

            parsed: list[list[str]] = []
            for rl in raw_rows:
                cells = [c.strip() for c in rl.strip("|").split("|")]
                parsed.append(cells)

            col_count = max(len(r) for r in parsed)

            table = doc.add_table(rows=len(parsed), cols=col_count)
            table.style = "Table Grid"

            for r_idx, row_data in enumerate(parsed):
                row = table.rows[r_idx]
                for c_idx in range(col_count):
                    cell = row.cells[c_idx]
                    cell_text = row_data[c_idx] if c_idx < len(row_data) else ""

                    # Clear auto-created empty paragraph and write content
                    p = cell.paragraphs[0]
                    p.clear()
                    _write_inline(p, cell_text)

                    if r_idx == 0:
                        # Header: dark blue bg, white bold text
                        _shade_cell(cell, "1F3A8A")
                        for run in p.runs:
                            run.bold = True
                            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    elif r_idx % 2 == 0:
                        _shade_cell(cell, "EEF2FF")

            doc.add_paragraph()
            continue

        # Bullet list
        if re.match(r"^[-*+]\s", stripped):
            p = doc.add_paragraph(style="List Bullet")
            _write_inline(p, stripped[2:])
            i += 1
            continue

        # Numbered list
        if re.match(r"^\d+\.\s", stripped):
            p = doc.add_paragraph(style="List Number")
            _write_inline(p, re.sub(r"^\d+\.\s+", "", stripped))
            i += 1
            continue

        # Blank line
        if not stripped:
            i += 1
            continue

        # Normal paragraph
        p = doc.add_paragraph()
        _write_inline(p, stripped)
        i += 1

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
