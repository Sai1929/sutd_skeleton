"""Build a formatted Word document (.docx) from the full RA JSON dict."""
import io
from datetime import date

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


# ── helpers ──────────────────────────────────────────────────────────────────

def _shade_cell(cell, hex_color: str) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _set_col_widths(table, widths_cm: list[float]) -> None:
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            if idx < len(widths_cm):
                cell.width = int(widths_cm[idx] * 360000)


def _header_row(table, labels: list[str], bg: str = "1F3A8A") -> None:
    row = table.rows[0]
    for idx, label in enumerate(labels):
        cell = row.cells[idx]
        p = cell.paragraphs[0]
        p.clear()
        run = p.add_run(label)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        _shade_cell(cell, bg)


def _cell_text(cell, text: str, size: int = 9, bold: bool = False, color: str | None = None) -> None:
    p = cell.paragraphs[0]
    p.clear()
    run = p.add_run(str(text))
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        r, g, b = int(color[0:2], 16), int(color[2:4], 16), int(color[4:6], 16)
        run.font.color.rgb = RGBColor(r, g, b)


def _risk_color(risk_str: str) -> str:
    r = risk_str.lower()
    if "very high" in r: return "B91C1C"
    if "high" in r:      return "D97706"
    if "medium" in r:    return "2563EB"
    return "16A34A"


def _section_heading(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x8A)


# ── main builder ─────────────────────────────────────────────────────────────

def json_to_docx(ra: dict, project_name: str = "Risk Assessment") -> bytes:
    doc = Document()

    for section in doc.sections:
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
        section.top_margin = Inches(0.9)
        section.bottom_margin = Inches(0.9)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10)

    # ── Cover / title ────────────────────────────────────────────────────────
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run(ra.get("project", project_name))
    title_run.bold = True
    title_run.font.size = Pt(18)
    title_run.font.color.rgb = RGBColor(0x0B, 0x12, 0x20)

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_p.add_run("RISK ASSESSMENT")
    sub_run.bold = True
    sub_run.font.size = Pt(12)
    sub_run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x8A)

    doc.add_paragraph()

    # ── Document info table ──────────────────────────────────────────────────
    info_table = doc.add_table(rows=4, cols=4)
    info_table.style = "Table Grid"
    info_fields = [
        ("Document No.",  ra.get("document_no", "RA-2024-001"), "Revision", ra.get("revision", "Rev 0")),
        ("Date",          ra.get("date", date.today().strftime("%d/%m/%Y")), "Approved By", ra.get("approved_by", "Site Manager")),
        ("Prepared By",   ra.get("prepared_by", "WSH Officer"), "Reviewed By", ra.get("reviewed_by", "Project Manager")),
        ("Scope",         ra.get("scope", ""), "", ""),
    ]
    for r_idx, (l1, v1, l2, v2) in enumerate(info_fields):
        cells = info_table.rows[r_idx].cells
        _cell_text(cells[0], l1, bold=True, size=9)
        _shade_cell(cells[0], "EEF2FF")
        _cell_text(cells[1], v1, size=9)
        _cell_text(cells[2], l2, bold=True, size=9)
        _shade_cell(cells[2], "EEF2FF")
        _cell_text(cells[3], v2, size=9)

    doc.add_paragraph()

    # ── Purpose ──────────────────────────────────────────────────────────────
    purpose = ra.get("purpose", "")
    if purpose:
        _section_heading(doc, "1. Purpose")
        doc.add_paragraph(purpose)

    # ── RA Table ─────────────────────────────────────────────────────────────
    rows = ra.get("rows", [])
    if rows:
        _section_heading(doc, "2. Risk Assessment")
        RA_COLS = [
            "Main Activity", "Sub-Activity", "Hazard", "Consequences",
            "L", "S", "Initial Risk", "Control Measures", "R-L", "R-S", "Residual Risk",
        ]
        ra_table = doc.add_table(rows=len(rows) + 1, cols=len(RA_COLS))
        ra_table.style = "Table Grid"
        _header_row(ra_table, RA_COLS)

        for r_idx, row in enumerate(rows):
            tr = ra_table.rows[r_idx + 1]
            vals = [
                row.get("main_activity", ""),
                row.get("sub_activity", ""),
                row.get("hazard", ""),
                row.get("consequences", ""),
                str(row.get("initial_l", "")),
                str(row.get("initial_s", "")),
                row.get("initial_risk", ""),
                row.get("control_measures", ""),
                str(row.get("residual_l", "")),
                str(row.get("residual_s", "")),
                row.get("residual_risk", ""),
            ]
            bg = "EEF2FF" if r_idx % 2 == 0 else "FFFFFF"
            for c_idx, val in enumerate(vals):
                cell = tr.cells[c_idx]
                p = cell.paragraphs[0]
                p.clear()
                run = p.add_run(val)
                run.font.size = Pt(8)
                _shade_cell(cell, bg)
                # Colour risk cells
                if c_idx in (6, 10) and val:
                    run.bold = True
                    col = _risk_color(val)
                    r, g, b = int(col[0:2], 16), int(col[2:4], 16), int(col[4:6], 16)
                    run.font.color.rgb = RGBColor(r, g, b)

        doc.add_paragraph()

    # ── Risk Matrix ───────────────────────────────────────────────────────────
    risk_matrix = ra.get("risk_matrix", {})
    if risk_matrix:
        _section_heading(doc, "3. Risk Matrix")
        note = risk_matrix.get("note", "")
        if note:
            p = doc.add_paragraph(note)
            p.runs[0].font.size = Pt(9)

        bands = risk_matrix.get("bands", [])
        if bands:
            band_table = doc.add_table(rows=len(bands) + 1, cols=3)
            band_table.style = "Table Grid"
            _header_row(band_table, ["Score Range", "Risk Level", "Required Action"])
            BAND_COLORS = {"Low": "16A34A", "Medium": "2563EB", "High": "D97706", "Very High": "B91C1C"}
            for r_idx, band in enumerate(bands):
                tr = band_table.rows[r_idx + 1]
                level = band.get("level", "")
                color = BAND_COLORS.get(level, "5A6272")
                _cell_text(tr.cells[0], band.get("range", ""), size=9)
                _cell_text(tr.cells[1], level, size=9, bold=True, color=color)
                _cell_text(tr.cells[2], band.get("action", ""), size=9)

        doc.add_paragraph()

    # ── Emergency Response ────────────────────────────────────────────────────
    emergency = ra.get("emergency_response", [])
    if emergency:
        _section_heading(doc, "4. Emergency Response")
        for item in emergency:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(item).font.size = Pt(9)
        doc.add_paragraph()

    # ── Chemical Note ─────────────────────────────────────────────────────────
    chem = ra.get("chemical_note", "")
    if chem:
        _section_heading(doc, "5. Chemical Hazards Note")
        p = doc.add_paragraph(chem)
        p.runs[0].font.size = Pt(9)
        doc.add_paragraph()

    # ── References ────────────────────────────────────────────────────────────
    references = ra.get("references", [])
    if references:
        _section_heading(doc, "6. References")
        for ref in references:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(ref).font.size = Pt(9)
        doc.add_paragraph()

    # ── Review Schedule ───────────────────────────────────────────────────────
    review = ra.get("review_schedule", "")
    if review:
        _section_heading(doc, "7. Review Schedule")
        p = doc.add_paragraph(review)
        p.runs[0].font.size = Pt(9)

    # ── Sign-off table ────────────────────────────────────────────────────────
    doc.add_paragraph()
    _section_heading(doc, "8. Sign-off")
    sign_table = doc.add_table(rows=2, cols=3)
    sign_table.style = "Table Grid"
    _header_row(sign_table, ["Prepared By", "Reviewed By", "Approved By"])
    sign_row = sign_table.rows[1]
    for cell, name in zip(sign_row.cells, [
        ra.get("prepared_by", "WSH Officer"),
        ra.get("reviewed_by", "Project Manager"),
        ra.get("approved_by", "Site Manager"),
    ]):
        p = cell.paragraphs[0]
        p.clear()
        p.add_run(f"\n\n\nName: {name}\nDate: ___________\nSignature: ___________").font.size = Pt(9)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
